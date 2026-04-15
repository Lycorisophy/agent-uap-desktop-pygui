"""
UAP 文档导入解析系统 - 文档解析器

提供 PDF、Word、Markdown 等格式的解析功能。
支持结构化内容提取和系统实体识别。
"""

import re
from abc import ABC, abstractmethod
from typing import List, Optional, Pattern
import json

from uap.document.models import (
    Document,
    DocumentType,
    ParsedContent,
    SystemEntity,
    EntityType,
)


class DocumentParser(ABC):
    """
    文档解析器抽象基类
    
    定义文档解析的接口。
    """
    
    # 数学方程模式
    EQUATION_PATTERNS: List[Pattern] = [
        # LaTeX 方程
        re.compile(r'\$\$(.+?)\$\$', re.DOTALL),
        re.compile(r'\$(.+?)\$'),
        # LaTeX 环境
        re.compile(r'\\begin\{(equation|align|gather|multline)\}(.+?)\\end\{\1\}', re.DOTALL),
        # ASCII 方程
        re.compile(r'([A-Za-z_][A-Za-z0-9_]*)\s*=\s*([^;]+)'),
        # 差分方程
        re.compile(r'd([A-Za-z_][A-Za-z0-9_]*)/dt\s*=\s*([^;]+)', re.IGNORECASE),
        re.compile(r'Δ([A-Za-z_][A-Za-z0-9_]*)\s*=\s*([^;]+)', re.IGNORECASE),
    ]
    
    # 变量模式
    VARIABLE_PATTERNS: List[Pattern] = [
        # x, y, t 等单字母
        re.compile(r'\b([x-yX-Y])\s*[:：]\s*([^,\n]+)'),
        # name = value 格式
        re.compile(r'\b([A-Za-z_][A-Za-z0-9_]*)\s*=\s*([\d\.\-]+)'),
        # 变量定义格式
        re.compile(r'变量\s*[:：]\s*([^\n]+)', re.IGNORECASE),
        re.compile(r'状态变量\s*[:：]\s*([^\n]+)', re.IGNORECASE),
    ]
    
    def __init__(self):
        self.parser_name = self.__class__.__name__
    
    @abstractmethod
    def parse(self, document: Document) -> ParsedContent:
        """
        解析文档
        
        Args:
            document: 文档对象
            
        Returns:
            解析后的内容
        """
        pass
    
    def extract_text(self, document: Document) -> str:
        """
        提取文本内容
        
        子类可重写此方法以提供自定义文本提取逻辑。
        """
        return document.content
    
    def extract_sections(self, text: str) -> List[dict]:
        """
        提取文档章节结构
        
        Args:
            text: 文档文本
            
        Returns:
            章节列表
        """
        sections = []
        
        # 按标题分割
        lines = text.split('\n')
        current_section = None
        
        for i, line in enumerate(lines):
            # 检测标题
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                title = heading_match.group(2).strip()
                
                if current_section:
                    sections.append(current_section)
                
                current_section = {
                    "level": level,
                    "title": title,
                    "content": "",
                    "line_start": i,
                    "line_end": i
                }
            elif current_section:
                current_section["content"] += line + "\n"
                current_section["line_end"] = i
        
        if current_section:
            sections.append(current_section)
        
        return sections
    
    def extract_equations(self, text: str) -> List[str]:
        """
        提取数学方程
        
        Args:
            text: 文档文本
            
        Returns:
            方程列表
        """
        equations = []
        
        for pattern in self.EQUATION_PATTERNS:
            matches = pattern.findall(text)
            for match in matches:
                equation = match.strip() if isinstance(match, str) else match[0].strip()
                if equation and len(equation) > 2:
                    equations.append(equation)
        
        # 去重
        seen = set()
        unique_equations = []
        for eq in equations:
            if eq not in seen:
                seen.add(eq)
                unique_equations.append(eq)
        
        return unique_equations
    
    def extract_entities(self, text: str) -> List[SystemEntity]:
        """
        提取系统实体
        
        子类可重写此方法以提供更精确的实体提取逻辑。
        """
        entities = []
        
        # 提取变量
        for pattern in self.VARIABLE_PATTERNS:
            matches = pattern.finditer(text)
            for match in matches:
                groups = match.groups()
                name = groups[0].strip()
                value = groups[1].strip() if len(groups) > 1 else None
                
                entity = SystemEntity(
                    entity_type=EntityType.VARIABLE,
                    name=name,
                    value=value,
                    description=f"从文本提取: {match.group(0)[:100]}",
                    source=f"位置: {match.start()}",
                    confidence=0.6
                )
                entities.append(entity)
        
        return entities
    
    def extract_tables(self, text: str) -> List[dict]:
        """
        提取表格数据
        
        Args:
            text: 文档文本
            
        Returns:
            表格列表
        """
        tables = []
        
        # Markdown 表格格式
        lines = text.split('\n')
        table_start = -1
        headers = []
        rows = []
        
        for i, line in enumerate(lines):
            # 检测表格分隔符
            if re.match(r'^\|[\s\-:|]+\|$', line):
                if table_start >= 0:
                    # 表格结束，提取数据
                    tables.append({
                        "headers": headers,
                        "rows": rows,
                        "line_start": table_start,
                        "line_end": i
                    })
                else:
                    # 解析表头
                    header_line = lines[i - 1] if i > 0 else ""
                    headers = [h.strip() for h in header_line.split('|') if h.strip()]
                    table_start = i - 1
                    rows = []
                continue
            
            # 检测表格行
            if '|' in line and table_start >= 0:
                cells = [c.strip() for c in line.split('|') if c.strip()]
                if cells and not cells[0].startswith('-'):
                    rows.append(cells)
        
        return tables
    
    def _create_parsed_content(
        self,
        document: Document,
        text: str,
        sections: List[dict] = None,
        entities: List[SystemEntity] = None,
        equations: List[str] = None,
        tables: List[dict] = None
    ) -> ParsedContent:
        """创建 ParsedContent 对象"""
        entities = entities or []
        equations = equations or []
        sections = sections or []
        tables = tables or []
        
        # 计算提取质量
        confidence = self._calculate_confidence(text, entities, equations)
        
        return ParsedContent(
            document=document,
            full_text=text,
            sections=sections,
            entities=entities,
            equations=equations,
            tables=tables,
            extraction_confidence=confidence,
            coverage=min(len(text) / 1000, 1.0)  # 简化的覆盖率计算
        )
    
    def _calculate_confidence(
        self,
        text: str,
        entities: List[SystemEntity],
        equations: List[str]
    ) -> float:
        """计算提取置信度"""
        score = 0.3  # 基础分
        
        # 有实体
        if len(entities) > 0:
            score += 0.2
        
        # 实体数量加成
        if len(entities) >= 3:
            score += 0.1
        
        # 有方程
        if len(equations) > 0:
            score += 0.2
        
        # 方程数量加成
        if len(equations) >= 2:
            score += 0.1
        
        # 内容长度
        if len(text) > 500:
            score += 0.1
        
        return min(score, 1.0)


class MarkdownParser(DocumentParser):
    """
    Markdown 文档解析器
    
    专门解析 Markdown 格式的文档。
    支持:
    - 标题和章节结构
    - 代码块和数学公式
    - 表格
    - 链接和图片
    - 列表
    """
    
    def __init__(self):
        super().__init__()
        self.parser_name = "MarkdownParser"
    
    def parse(self, document: Document) -> ParsedContent:
        """解析 Markdown 文档"""
        if not document.is_markdown and document.doc_type != DocumentType.MARKDOWN:
            raise ValueError(f"Expected markdown document, got {document.doc_type}")
        
        text = self.extract_text(document)
        sections = self.extract_sections(text)
        equations = self.extract_equations(text)
        tables = self.extract_tables(text)
        
        # Markdown 优化的实体提取
        entities = self._extract_markdown_entities(text)
        
        return self._create_parsed_content(
            document, text, sections, entities, equations, tables
        )
    
    def _extract_markdown_entities(self, text: str) -> List[SystemEntity]:
        """提取 Markdown 中的系统实体"""
        entities = []
        
        # 提取代码块中的变量定义
        code_blocks = re.findall(r'```[\s\S]*?```', text)
        for block in code_blocks:
            block_entities = self.extract_entities(block)
            for entity in block_entities:
                entity.source = "代码块"
                entity.confidence = 0.7
                entities.append(entity)
        
        # 提取 Markdown 表格中的变量
        table_entities = self._extract_table_variables(text)
        entities.extend(table_entities)
        
        # 提取列表中的变量定义
        list_entities = self._extract_list_variables(text)
        entities.extend(list_entities)
        
        return entities
    
    def _extract_table_variables(self, text: str) -> List[SystemEntity]:
        """从表格中提取变量"""
        entities = []
        tables = self.extract_tables(text)
        
        for table in tables:
            headers = table.get("headers", [])
            
            # 检测参数表
            param_keywords = ["参数", "parameter", "变量", "variable", "符号", "symbol"]
            is_param_table = any(kw in ' '.join(headers).lower() for kw in param_keywords)
            
            if is_param_table:
                for row in table.get("rows", []):
                    if len(row) >= 2:
                        name = row[0].strip()
                        value = row[1].strip()
                        
                        if name and value:
                            entity = SystemEntity(
                                entity_type=EntityType.PARAMETER,
                                name=name,
                                value=value,
                                description="从参数表提取",
                                source=f"表格: {headers}",
                                confidence=0.8
                            )
                            entities.append(entity)
        
        return entities
    
    def _extract_list_variables(self, text: str) -> List[SystemEntity]:
        """从列表中提取变量"""
        entities = []
        
        # 检测变量列表
        list_patterns = [
            r'变量\s*[：:]\s*([^\n]+)',
            r'参数\s*[：:]\s*([^\n]+)',
            r'常数\s*[：:]\s*([^\n]+)',
        ]
        
        for pattern in list_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                content = match.group(1)
                # 解析逗号分隔的列表
                items = re.split(r'[,，、]', content)
                for item in items:
                    item = item.strip()
                    if item and len(item) < 50:
                        # 尝试提取名称和值
                        name_value = re.match(
                            r'([A-Za-z_][A-Za-z0-9_]*)\s*[=＝]\s*([\d\.\-]+)', item
                        )
                        if name_value:
                            entity = SystemEntity(
                                entity_type=EntityType.VARIABLE,
                                name=name_value.group(1),
                                value=name_value.group(2),
                                description="从列表提取",
                                source="变量列表",
                                confidence=0.7,
                            )
                        else:
                            entity = SystemEntity(
                                entity_type=EntityType.VARIABLE,
                                name=item,
                                description="从列表提取",
                                source="变量列表",
                                confidence=0.5,
                            )
                        entities.append(entity)
        
        return entities


class PdfParser(DocumentParser):
    """
    PDF 文档解析器
    
    解析 PDF 格式的文档。
    需要安装 PyMuPDF (fitz) 或 pdfplumber。
    """
    
    def __init__(self):
        super().__init__()
        self.parser_name = "PdfParser"
        self._pdf_engine = None
        self._setup_engine()
    
    def _setup_engine(self):
        """设置 PDF 解析引擎"""
        # 尝试 PyMuPDF
        try:
            import fitz
            self._pdf_engine = "pymupdf"
            return
        except ImportError:
            pass
        
        # 尝试 pdfplumber
        try:
            import pdfplumber
            self._pdf_engine = "pdfplumber"
            return
        except ImportError:
            pass
        
        self._pdf_engine = None
    
    def parse(self, document: Document) -> ParsedContent:
        """解析 PDF 文档"""
        if not document.is_pdf and document.doc_type != DocumentType.PDF:
            raise ValueError(f"Expected PDF document, got {document.doc_type}")
        
        if self._pdf_engine is None:
            raise ImportError(
                "PDF parsing requires PyMuPDF or pdfplumber. "
                "Install with: pip install pymupdf  OR  pip install pdfplumber"
            )
        
        # 提取文本
        text = self.extract_text(document)
        
        # 提取章节（基于标题）
        sections = self._extract_pdf_sections(text)
        
        # 提取方程和实体
        equations = self.extract_equations(text)
        entities = self.extract_entities(text)
        
        # 提取表格
        tables = self._extract_pdf_tables(document)
        
        return self._create_parsed_content(
            document, text, sections, entities, equations, tables
        )
    
    def extract_text(self, document: Document) -> str:
        """从 PDF 提取文本"""
        if self._pdf_engine == "pymupdf":
            return self._extract_pymupdf(document)
        elif self._pdf_engine == "pdfplumber":
            return self._extract_pdfplumber(document)
        else:
            return document.content
    
    def _extract_pymupdf(self, document: Document) -> str:
        """使用 PyMuPDF 提取文本"""
        import fitz
        
        text_parts = []
        
        try:
            # 尝试作为文件路径
            if document.file_path:
                doc = fitz.open(document.file_path)
            else:
                # 使用内存中的数据
                import io
                doc = fitz.open(stream=document.content.encode(), filetype="pdf")
            
            for page_num, page in enumerate(doc):
                text = page.get_text()
                text_parts.append(f"--- 第 {page_num + 1} 页 ---\n{text}")
            
            doc.close()
            
        except Exception as e:
            print(f"PyMuPDF extraction failed: {e}")
            return document.content
        
        return "\n\n".join(text_parts)
    
    def _extract_pdfplumber(self, document: Document) -> str:
        """使用 pdfplumber 提取文本"""
        import pdfplumber
        
        text_parts = []
        
        try:
            with pdfplumber.open(document.file_path or document.content) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    text_parts.append(f"--- 第 {page_num + 1} 页 ---\n{text}")
                    
        except Exception as e:
            print(f"pdfplumber extraction failed: {e}")
            return document.content
        
        return "\n\n".join(text_parts)
    
    def _extract_pdf_sections(self, text: str) -> List[dict]:
        """从 PDF 文本提取章节"""
        sections = []
        
        # PDF 通常没有 Markdown 标题，使用其他方法
        # 1. 大号字体行（标题）
        lines = text.split('\n')
        
        for i, line in enumerate(lines):
            stripped = line.strip()
            
            # 检测可能的标题行
            # - 全大写
            # - 短行 + 数字结尾
            # - 特定关键词
            is_heading = (
                (stripped.isupper() and len(stripped) < 100) or
                re.match(r'.+\d+\.\s*$', stripped) or
                re.match(r'^(第|Chapter|Section)\s*\d+', stripped, re.IGNORECASE)
            )
            
            if is_heading and len(stripped) > 3:
                sections.append({
                    "title": stripped,
                    "level": 1,
                    "content": "",
                    "line_start": i
                })
        
        return sections
    
    def _extract_pdf_tables(self, document: Document) -> List[dict]:
        """从 PDF 提取表格"""
        tables = []
        
        if self._pdf_engine == "pdfplumber":
            try:
                import pdfplumber
                
                with pdfplumber.open(document.file_path or document.content) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        page_tables = page.extract_tables()
                        
                        for table in page_tables:
                            if table and len(table) > 1:
                                tables.append({
                                    "headers": table[0] if table else [],
                                    "rows": table[1:] if len(table) > 1 else [],
                                    "page": page_num + 1
                                })
                                
            except Exception as e:
                print(f"Table extraction failed: {e}")
        
        return tables


class WordParser(DocumentParser):
    """
    Word 文档解析器
    
    解析 Word 格式的文档 (.docx)。
    需要安装 python-docx。
    """
    
    def __init__(self):
        super().__init__()
        self.parser_name = "WordParser"
    
    def parse(self, document: Document) -> ParsedContent:
        """解析 Word 文档"""
        if not document.is_word and document.doc_type != DocumentType.WORD:
            raise ValueError(f"Expected Word document, got {document.doc_type}")
        
        text = self.extract_text(document)
        sections = self._extract_word_sections(document)
        equations = self.extract_equations(text)
        entities = self._extract_word_entities(document)
        tables = self._extract_word_tables(document)
        
        return self._create_parsed_content(
            document, text, sections, entities, equations, tables
        )
    
    def extract_text(self, document: Document) -> str:
        """从 Word 提取文本"""
        try:
            import docx
        except ImportError:
            raise ImportError(
                "Word parsing requires python-docx. "
                "Install with: pip install python-docx"
            )
        
        text_parts = []
        
        try:
            doc = docx.Document(document.file_path)
            
            # 提取段落
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # 提取表格文本
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    if any(row_text):
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    text_parts.append("\n".join(table_text))
            
        except Exception as e:
            print(f"Word extraction failed: {e}")
            return document.content
        
        return "\n\n".join(text_parts)
    
    def _extract_word_sections(self, document: Document) -> List[dict]:
        """从 Word 提取章节"""
        sections = []
        
        try:
            import docx
            
            doc = docx.Document(document.file_path)
            
            for para in doc.paragraphs:
                style = para.style.name if para.style else ""
                
                # 检测标题样式
                if "Heading" in style or para.text.strip().startswith(("#", "第", "Chapter", "Section")):
                    sections.append({
                        "title": para.text.strip(),
                        "level": self._heading_level(style),
                        "content": "",
                        "style": style
                    })
                    
        except Exception as e:
            print(f"Section extraction failed: {e}")
        
        return sections
    
    def _heading_level(self, style: str) -> int:
        """从样式名获取标题级别"""
        match = re.search(r'(\d+)', style)
        if match:
            return int(match.group(1))
        return 1
    
    def _extract_word_entities(self, document: Document) -> List[SystemEntity]:
        """从 Word 提取实体"""
        entities = []
        
        try:
            import docx
            
            doc = docx.Document(document.file_path)
            
            # 从表格中提取参数
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    
                    if len(cells) >= 2:
                        # 检测参数列
                        if any(kw in cells[0].lower() for kw in ["参数", "parameter", "变量", "variable"]):
                            name = cells[0]
                            value = cells[1] if len(cells) > 1 else None
                            
                            entity = SystemEntity(
                                entity_type=EntityType.PARAMETER,
                                name=name,
                                value=value,
                                description="从Word表格提取",
                                source="参数表",
                                confidence=0.8
                            )
                            entities.append(entity)
            
            # 从段落中提取
            full_text = self.extract_text(document)
            paragraph_entities = self.extract_entities(full_text)
            entities.extend(paragraph_entities)
            
        except Exception as e:
            print(f"Entity extraction failed: {e}")
        
        return entities
    
    def _extract_word_tables(self, document: Document) -> List[dict]:
        """从 Word 提取表格"""
        tables = []
        
        try:
            import docx
            
            doc = docx.Document(document.file_path)
            
            for table in doc.tables:
                headers = []
                rows = []
                
                for row_idx, row in enumerate(table.rows):
                    cells = [cell.text.strip() for cell in row.cells]
                    
                    if row_idx == 0:
                        headers = cells
                    else:
                        rows.append(cells)
                
                tables.append({
                    "headers": headers,
                    "rows": rows
                })
                
        except Exception as e:
            print(f"Table extraction failed: {e}")
        
        return tables


def create_parser(doc_type: DocumentType = None) -> DocumentParser:
    """
    创建文档解析器
    
    Args:
        doc_type: 文档类型
        
    Returns:
        解析器实例
    """
    if doc_type == DocumentType.MARKDOWN:
        return MarkdownParser()
    elif doc_type == DocumentType.PDF:
        return PdfParser()
    elif doc_type == DocumentType.WORD:
        return WordParser()
    else:
        # 默认返回 Markdown 解析器
        return MarkdownParser()


def create_parser_for_file(file_path: str) -> DocumentParser:
    """
    根据文件路径创建解析器
    
    Args:
        file_path: 文件路径
        
    Returns:
        解析器实例
    """
    from pathlib import Path
    
    ext = Path(file_path).suffix.lower()
    
    if ext in [".md", ".markdown"]:
        return MarkdownParser()
    elif ext == ".pdf":
        return PdfParser()
    elif ext in [".docx", ".doc"]:
        return WordParser()
    else:
        return MarkdownParser()
